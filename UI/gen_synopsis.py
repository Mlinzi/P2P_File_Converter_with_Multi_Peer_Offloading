from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Distributed P2P File Conversion Network')
run.bold = True
run.font.size = Pt(16)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('Computer Networks Mini Project – Synopsis')
run2.bold = True
run2.font.size = Pt(12)

ids = doc.add_paragraph()
ids.alignment = WD_ALIGN_PARAGRAPH.CENTER
ids.add_run('PES2UG24AM047 | PES2UG24AM032 | PES2UG24AM037').font.size = Pt(11)

doc.add_paragraph()

def section(num, title_text):
    p = doc.add_paragraph()
    run = p.add_run(f'{num}. {title_text}')
    run.bold = True
    run.font.size = Pt(12)

# 1
section(1, 'Abstract – Problem Statement and Scope')
doc.add_paragraph(
    'This project designs and implements a Distributed Peer-to-Peer (P2P) File Conversion Network '
    'where every participating node is architecturally equal — capable of both requesting and fulfilling '
    'file conversion jobs. Unlike centralized architectures, there is no coordinator server, no load balancer, '
    'and no shared task queue. Peers discover each other autonomously using mDNS (Multicast DNS) on a local '
    'network and negotiate conversion jobs directly over secure SSL/TLS TCP connections.\n\n'
    'The system addresses real-world challenges of scalability, fault tolerance, and decentralized resource '
    'utilization by distributing conversion workload across all available peers. All file transfers and control '
    'messages occur exclusively over raw TCP sockets secured with SSL/TLS, satisfying the mandatory security '
    'requirements. The project demonstrates core networking concepts including socket programming, P2P protocol '
    'design, decentralized job scheduling, concurrency, and performance evaluation under realistic conditions.'
)

# 2
section(2, 'Architecture Design')
doc.add_paragraph(
    'The system follows a Pure Peer-to-Peer Mesh Architecture. All nodes are functionally identical — '
    'each peer runs the same software stack and can simultaneously act as a requester and a converter.'
)
items = [
    'Peer Discovery Module – Uses mDNS (zeroconf) to announce presence and discover other peers on the LAN automatically, with no central tracker or bootstrap server required.',
    'Job Negotiation Protocol – When a peer needs a conversion, it broadcasts a job announcement to known peers. Available peers respond with their current load and capabilities. The requester selects the best peer.',
    'SSL/TLS Socket Layer – All peer-to-peer communication (job negotiation, file transfer, status updates) occurs over raw TCP sockets wrapped with SSL/TLS using self-signed certificates.',
    'Conversion Engine – Each peer runs FFmpeg (media) and LibreOffice CLI (documents) locally to execute conversion tasks.',
    'Metrics & Dashboard – Each peer hosts a local Flask dashboard exposing real-time performance metrics via a web UI.',
]
for item in items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

# 3
section(3, 'Architectural Workflow')
steps = [
    'Peer starts up and announces itself on the LAN via mDNS. It discovers other active peers automatically.',
    'Peer A needs to convert a file (e.g., MP4 to MP3). It broadcasts a JobRequest message to all known peers over SSL/TLS TCP.',
    'Available peers respond with a JobOffer containing their current CPU load and queue length.',
    'Peer A selects the best available peer (e.g., Peer B with lowest load) and establishes a direct SSL/TLS TCP connection.',
    'Peer A sends the file in binary chunks over the secure socket to Peer B.',
    'Peer B executes the conversion locally using FFmpeg or LibreOffice CLI.',
    'Peer B sends the converted file back to Peer A over the same SSL/TLS channel.',
    'Both peers update their local metrics (latency, throughput, CPU usage) visible on their dashboards.',
    'If Peer B disconnects mid-transfer, Peer A retries with the next available peer — no central coordinator needed.',
]
for step in steps:
    p = doc.add_paragraph(style='List Number')
    p.add_run(step)

# 4
section(4, 'Relevance of Architecture Design and Choice')
points = [
    ('True Decentralization', 'No single point of failure. Any peer can leave or join without disrupting the network.'),
    ('Scalability', 'Adding a new peer automatically increases the conversion capacity — no reconfiguration needed.'),
    ('Security', 'SSL/TLS on every peer-to-peer channel ensures confidentiality and integrity, satisfying the mandatory requirement.'),
    ('Fault Tolerance', 'If a peer fails mid-conversion, the requester renegotiates with another available peer.'),
    ('Concurrency', 'Each peer handles multiple simultaneous connections using threading, enabling parallel conversions across the network.'),
    ('Real-world Applicability', 'Mirrors architectures used in BitTorrent, Syncthing, and decentralized compute networks.'),
]
for t, desc in points:
    p = doc.add_paragraph(style='List Bullet')
    run_b = p.add_run(t + ': ')
    run_b.bold = True
    p.add_run(desc)

# 5
section(5, 'Set of Objectives')
objectives = [
    'Implement low-level TCP socket communication using Python\'s socket library for all peer-to-peer data exchange.',
    'Enforce SSL/TLS on all peer connections using self-signed certificates.',
    'Implement mDNS-based peer discovery — no manual IP configuration required.',
    'Design and implement a custom P2P job negotiation protocol over raw sockets.',
    'Support binary file transfer over the socket layer between peers.',
    'Handle multiple concurrent peer connections using threading.',
    'Execute format conversion (audio, video, document, image) on each peer node.',
    'Measure and visualize performance metrics in real-time via a per-peer web dashboard.',
    'Demonstrate fault tolerance by handling peer disconnection and job retry.',
]
for obj in objectives:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(obj)

# 6
section(6, 'Functional Requirements')
reqs = [
    'Automatic peer discovery on LAN using mDNS (no manual IP entry).',
    'Job broadcast and negotiation protocol between peers over SSL/TLS TCP.',
    'Binary file upload and download directly between peers at the socket level.',
    'Local conversion engine on each peer (FFmpeg for media, LibreOffice CLI for documents).',
    'Concurrent handling of multiple inbound and outbound connections per peer.',
    'SSL/TLS certificate generation and management for each peer node.',
    'Real-time performance metrics dashboard (Flask) on each peer.',
    'Peer failure detection and automatic job retry with alternate peer.',
]
for req in reqs:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(req)

# 7
section(7, 'Software and Hardware Requirements')
doc.add_paragraph()

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Component'
hdr[1].text = 'Technology / Detail'
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True

rows_data = [
    ('Socket Layer', 'Python (socket, ssl, threading) – raw TCP + TLS'),
    ('Peer Discovery', 'zeroconf / python-zeroconf (mDNS)'),
    ('Conversion Tools', 'FFmpeg (audio/video), LibreOffice CLI (documents/images)'),
    ('Dashboard UI', 'Flask + Chart.js (per-peer local web dashboard)'),
    ('Security', 'SSL/TLS with self-signed certificates (openssl)'),
    ('Protocol', 'Custom JSON-based control messages over TCP'),
    ('Hardware', 'Any 2+ laptops on same LAN/hotspot. Min 4 GB RAM per node.'),
]
for comp, tech in rows_data:
    row = table.add_row().cells
    row[0].text = comp
    row[1].text = tech

doc.add_paragraph()

# 8
section(8, 'Parameters to Evaluate Performance')
doc.add_paragraph()

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Parameter'
hdr2[1].text = 'Description'
hdr2[2].text = 'Measurement Method'
for cell in hdr2:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True

perf_data = [
    ('Conversion Latency', 'Time from file send to converted file received', 'Timestamped at socket send/receive'),
    ('Throughput', 'Number of files successfully converted per unit time', 'Counter over time window'),
    ('CPU Utilization', 'Processing resource usage on converting peer', 'psutil on each peer node'),
    ('Bandwidth Usage', 'Network data consumed per conversion', 'Bytes sent/received at socket level'),
    ('Peer Scalability', 'Latency and throughput change as peer count increases', 'Benchmark with 2, 3, 4 peers'),
    ('Fault Tolerance', 'System recovery time on peer failure mid-transfer', 'Kill peer, measure retry time'),
    ('File Size vs Latency', 'Conversion time as function of input file size', 'Test with 1 MB, 10 MB, 100 MB'),
    ('TLS Overhead', 'Extra latency from SSL/TLS handshake vs plain TCP', 'Compare TLS vs raw TCP transfer time'),
]
for param, desc, method in perf_data:
    row = table2.add_row().cells
    row[0].text = param
    row[1].text = desc
    row[2].text = method

doc.add_paragraph()

# 9
section(9, 'Expected Outcome')
outcomes = [
    'A fully functional P2P file conversion network with SSL/TLS-secured TCP communication between peers.',
    'Automatic peer discovery on LAN — no manual configuration.',
    'Binary file transfer working reliably at the socket level between any two peers.',
    'Demonstrated reduction in conversion latency as more peers join the network.',
    'Real-time performance dashboard visible on each peer node.',
    'Empirical analysis of conversion time versus file size.',
    'Quantified TLS overhead relative to plain TCP.',
    'Demonstrated fault tolerance with automatic job retry on peer failure.',
]
for outcome in outcomes:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(outcome)

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run('Languages: Python | Protocol: TCP with SSL/TLS | Architecture: Pure P2P Mesh | Discovery: mDNS')
r.italic = True

doc.save('CN_Mini_Project_Synopsis_Team_18_P2P.docx')
print('Done')
